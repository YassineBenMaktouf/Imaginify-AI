from flask import Flask, render_template, request, jsonify , send_file
import os
import requests
from dotenv import load_dotenv
import time
import logging
from requests.exceptions import HTTPError, Timeout, RequestException
from flask import jsonify, request
from docx.shared import Inches
import io
from docx import Document

load_dotenv()

app = Flask(__name__)
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
HUGGINGFACE_API_KEY = os.getenv('HUGGINGFACE_API_KEY')
DICTIONARY_API_KEY = os.getenv('DICTIONARY_API_KEY')
OPENAI_API_URL = "https://api.openai.com/v1/completions"
HUGGINGFACE_API_URL = "https://api-inference.huggingface.co/models/runwayml/stable-diffusion-v1-5"

logging.basicConfig(level=logging.INFO)

def generate_text(prompt, model="gpt-3.5-turbo-instruct", max_tokens=300):
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}"}
    data = {
        "model": model,
        "prompt": prompt,
        "max_tokens": max_tokens,
        "temperature": 0.7,
        "n": 1,
        "stop": None
    }
    try:
        response = requests.post(OPENAI_API_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json()['choices'][0]['text'].strip()
    except HTTPError as http_err:
        logging.error(f'HTTP error occurred: {http_err}')
    except Exception as err:
        logging.error(f'Other error occurred: {err}')
    return None

def query_image_generation(prompt, retry_limit=3, timeout=10):
    headers = {"Authorization": f"Bearer {HUGGINGFACE_API_KEY}"}
    data = {"inputs": prompt}
    for attempt in range(retry_limit):
        try:
            response = requests.post(HUGGINGFACE_API_URL, headers=headers, json=data, timeout=timeout)
            response.raise_for_status()
            return response.content
        except Timeout:
            logging.warning('The request timed out, attempting retry...')
        except HTTPError as http_err:
            if response.status_code == 503 and attempt < retry_limit - 1:
                logging.warning('Service unavailable, retrying...')
            else:
                logging.error(f'HTTP error occurred: {http_err}')
                break
        except RequestException as req_err:
            logging.error(f'Request error occurred: {req_err}')
            break
        time.sleep((2 ** attempt) * 3)
    return None

def save_image(image_data, folder="static/generated_images"):
    if not os.path.exists(folder):
        os.makedirs(folder)
    timestamp = int(time.time())
    filename = f"image_{timestamp}.png"
    file_path = os.path.join(folder, filename)
    rel_path = os.path.join('generated_images', filename)
    with open(file_path, 'wb') as f:
        f.write(image_data)
    return rel_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate-story', methods=['POST'])
def generate_story():
    data = request.get_json()
    user_prompt = data['prompt']
    lines = data.get('lines', 10)
    story_prompt = f"Write a short story about {user_prompt} in {lines} lines."
    story = generate_text(story_prompt)
    if story:
        image_data = query_image_generation(story)
        if image_data:
            image_filename = save_image(image_data)
            return jsonify({"story": story, "image_filename": image_filename})
        else:
            return jsonify({"story": story})
    else:
        return jsonify({"error": "Failed to generate the story."}), 500


@app.route('/generate-continuation', methods=['POST'])
def generate_continuation():
    data = request.get_json()
    logging.info(f"Received data: {data}")
    index = data.get('index')
    full_story = data['story']
    lines = data.get('lines', 10)
    
    if index is not None and 0 <= index < len(full_story.split('\n\n')):
        story_parts = full_story.split('\n\n')
        continuation_prompt = f"Continue the following story part in {lines} lines:\n\n{story_parts[index]}"
    else:
        continuation_prompt = f"Continue the following story in {lines} lines:\n\n{full_story}"

    continuation = generate_text(continuation_prompt)
    if continuation:
        image_data = query_image_generation(continuation)
        if image_data:
            image_filename = save_image(image_data)
            response_data = {"continuation": continuation, "image_filename": image_filename}
        else:
            response_data = {"continuation": continuation}
        if index is not None:
            response_data['index'] = index

        return jsonify(response_data)
    else:
        return jsonify({"error": "Failed to generate continuation text."}), 500

@app.route('/generate-titles', methods=['POST'])
def generate_titles():
    data = request.get_json()
    story_text = data['story']
    titles_prompt = f"Based on the following story, suggest three titles:\n\n{story_text}\n\nTitles:"

    titles = generate_text(titles_prompt, max_tokens=200)
    
    if titles:
        titles_list = [title.strip() for title in titles.split('\n') if title.strip()]
        titles_list = titles_list[:3]
        return jsonify({"titles": titles_list})
    else:
        return jsonify({"error": "Failed to generate titles."}), 500
def generate_text(prompt, max_tokens):
    pass
DICTIONARY_API_KEY = os.getenv('DICTIONARY_API_KEY')

@app.route('/lookup-word', methods=['GET'])
def lookup_word():
    word = request.args.get('word', '').strip()
    if not word:
        return jsonify({"error": "No word provided."}), 400

    api_url = f'https://api.dictionaryapi.dev/api/v2/entries/en/{word}'

    try:
        response = requests.get(api_url, headers={'Authorization': f'Bearer {DICTIONARY_API_KEY}'})
        response.raise_for_status()
        definitions = response.json()

        formatted_definitions = [{"definition": definition.get("meanings", [{}])[0].get("definitions", [{}])[0].get("definition", "")} for definition in definitions]

        return jsonify(formatted_definitions)
    except requests.RequestException as e:
        print(f"API request failed: {e}")
        return jsonify({"error": "Failed to fetch the definition"}), 500

@app.route('/download-story', methods=['POST'])
def download_story():
    data = request.json
    doc = Document()

    for part in data.get('storyParts', []):
        doc.add_paragraph(part['text'])
        for image_url in part['images']:
            try:
                response = requests.get(image_url)
                if response.status_code == 200:
                    image_bytes = io.BytesIO(response.content)
                    doc.add_picture(image_bytes, width=Inches(2))
            except requests.RequestException:
                print(f"Failed to download image from {image_url}")

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)

    return send_file(f, as_attachment=True, download_name='YourStory.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)

